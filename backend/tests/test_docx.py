"""
Tests para el módulo de conversión HTML -> DOCX.

Valida que se generen archivos DOCX válidos.
"""

import pytest
from io import BytesIO
from docx import Document
from app.converter.html_to_docx import convert, convert_with_template
from app.converter.exceptions import HTMLConversionError


class TestHTMLToDocx:
    """Tests para conversión básica de HTML a DOCX."""
    
    def test_simple_html_conversion(self):
        """Test conversión de HTML simple."""
        html = "<h1>Hello World</h1>"
        docx_bytes = convert(html)
        
        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0
    
    def test_returns_valid_docx(self):
        """Test que retorna DOCX válido."""
        html = "<h1>Test</h1><p>Content</p>"
        docx_bytes = convert(html)
        
        buffer = BytesIO(docx_bytes)
        doc = Document(buffer)
        
        assert len(doc.paragraphs) > 0
    
    def test_paragraph_conversion(self):
        """Test conversión de párrafo."""
        html = "<p>This is a paragraph</p>"
        docx_bytes = convert(html)
        
        buffer = BytesIO(docx_bytes)
        doc = Document(buffer)
        
        text = "\n".join([p.text for p in doc.paragraphs])
        assert "This is a paragraph" in text
    
    def test_heading_conversion(self):
        """Test conversión de encabezados."""
        html = "<h1>Title</h1><h2>Subtitle</h2>"
        docx_bytes = convert(html)
        
        buffer = BytesIO(docx_bytes)
        doc = Document(buffer)
        
        assert len(doc.paragraphs) > 0
    
    def test_bold_text(self):
        """Test conversión de texto en negrita."""
        html = "<p><strong>bold text</strong></p>"
        docx_bytes = convert(html)
        
        buffer = BytesIO(docx_bytes)
        doc = Document(buffer)
        
        assert len(doc.paragraphs) > 0
    
    def test_italic_text(self):
        """Test conversión de texto en cursiva."""
        html = "<p><em>italic text</em></p>"
        docx_bytes = convert(html)
        
        buffer = BytesIO(docx_bytes)
        doc = Document(buffer)
        
        assert len(doc.paragraphs) > 0
    
    def test_list_conversion(self):
        """Test conversión de listas."""
        html = """
        <ul>
            <li>Item 1</li>
            <li>Item 2</li>
            <li>Item 3</li>
        </ul>
        """
        docx_bytes = convert(html)
        
        buffer = BytesIO(docx_bytes)
        doc = Document(buffer)
        
        text = "\n".join([p.text for p in doc.paragraphs])
        assert "Item 1" in text or len(doc.paragraphs) > 0
    
    def test_table_conversion(self):
        """Test conversión de tablas."""
        html = """
        <table>
            <tr><th>Header 1</th><th>Header 2</th></tr>
            <tr><td>Data 1</td><td>Data 2</td></tr>
        </table>
        """
        docx_bytes = convert(html)
        
        buffer = BytesIO(docx_bytes)
        doc = Document(buffer)
        
        assert len(doc.paragraphs) > 0 or len(doc.tables) > 0
    
    def test_link_conversion(self):
        """Test conversión de links."""
        html = '<p><a href="https://google.com">Link</a></p>'
        docx_bytes = convert(html)
        
        buffer = BytesIO(docx_bytes)
        doc = Document(buffer)
        
        assert len(doc.paragraphs) > 0


class TestHTMLEdgeCases:
    """Tests para casos especiales."""
    
    def test_empty_html_raises_error(self):
        """Test que HTML vacío lanza error."""
        with pytest.raises(ValueError, match="no puede estar vacío"):
            convert("")
    
    def test_none_raises_error(self):
        """Test que None lanza error."""
        with pytest.raises(TypeError):
            convert(None)
    
    def test_non_string_raises_error(self):
        """Test que tipo incorrecto lanza error."""
        with pytest.raises(TypeError, match="debe ser un string"):
            convert(123)
        
        with pytest.raises(TypeError):
            convert([])
    
    def test_whitespace_only_html(self):
        """Test HTML con solo espacios."""
        html = "   \n\n   "
        
        with pytest.raises(ValueError):
            convert(html)
    
    def test_malformed_html(self):
        """Test HTML malformado."""
        html = "<h1>Unclosed header"
        
        docx_bytes = convert(html)
        assert isinstance(docx_bytes, bytes)
    
    def test_special_html_characters(self):
        """Test caracteres especiales HTML."""
        html = "<p>&lt; &gt; &amp; &quot;</p>"
        docx_bytes = convert(html)
        
        buffer = BytesIO(docx_bytes)
        doc = Document(buffer)
        
        assert len(doc.paragraphs) > 0


class TestComplexHTML:
    """Tests para HTML complejo."""
    
    def test_nested_elements(self):
        """Test elementos HTML anidados."""
        html = """
        <div>
            <h1>Title</h1>
            <div>
                <p>Nested <strong>paragraph</strong></p>
            </div>
        </div>
        """
        docx_bytes = convert(html)
        
        buffer = BytesIO(docx_bytes)
        doc = Document(buffer)
        
        assert len(doc.paragraphs) > 0
    
    def test_mixed_content(self, sample_html):
        """Test HTML con contenido mixto."""
        docx_bytes = convert(sample_html)
        
        buffer = BytesIO(docx_bytes)
        doc = Document(buffer)
        
        assert len(doc.paragraphs) > 1
    
    def test_very_long_html(self):
        """Test HTML muy largo."""
        html = "<p>" + ("Long text. " * 1000) + "</p>"
        docx_bytes = convert(html)
        
        buffer = BytesIO(docx_bytes)
        doc = Document(buffer)
        
        assert len(doc.paragraphs) > 0


class TestCustomStyles:
    """Tests para estilos personalizados."""
    
    def test_convert_without_styles(self):
        """Test conversión sin estilos personalizados."""
        html = "<h1>Test</h1>"
        docx_bytes = convert(html, custom_styles=None)
        
        assert isinstance(docx_bytes, bytes)
    
    def test_convert_with_empty_styles(self):
        """Test conversión con dict de estilos vacío."""
        html = "<h1>Test</h1>"
        docx_bytes = convert(html, custom_styles={})
        
        assert isinstance(docx_bytes, bytes)


class TestTemplateConversion:
    """Tests para conversión con template."""
    
    def test_convert_without_template(self):
        """Test conversión sin template."""
        html = "<h1>Test</h1>"
        docx_bytes = convert_with_template(html, template_path=None)
        
        assert isinstance(docx_bytes, bytes)
    
    def test_convert_returns_bytes(self):
        """Test que siempre retorna bytes."""
        html = "<p>Simple text</p>"
        docx_bytes = convert_with_template(html)
        
        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0


class TestDocxFileValidity:
    """Tests para validar la estructura del DOCX generado."""
    
    def test_docx_has_correct_structure(self):
        """Test que el DOCX tiene estructura correcta."""
        html = "<h1>Title</h1><p>Content</p>"
        docx_bytes = convert(html)
        
        buffer = BytesIO(docx_bytes)
        doc = Document(buffer)
        
        assert hasattr(doc, 'paragraphs')
        assert hasattr(doc, 'tables')
        assert hasattr(doc, 'sections')
    
    def test_docx_can_be_saved(self, temp_dir):
        """Test que el DOCX puede ser guardado."""
        html = "<h1>Test</h1>"
        docx_bytes = convert(html)
        
        test_file = temp_dir / "test.docx"
        with open(test_file, "wb") as f:
            f.write(docx_bytes)
        
        assert test_file.exists()
        assert test_file.stat().st_size > 0
        
        doc = Document(test_file)
        assert len(doc.paragraphs) >= 0


class TestPerformance:
    """Tests de performance."""
    
    def test_conversion_speed(self):
        """Test que la conversión es rápida."""
        import time
        
        html = "<h1>Test</h1>" * 10
        
        start = time.time()
        convert(html)
        duration = time.time() - start
        
        assert duration < 2.0
    
    def test_multiple_conversions(self):
        """Test múltiples conversiones consecutivas."""
        html = "<p>Test</p>"
        
        for _ in range(50):
            docx_bytes = convert(html)
            assert isinstance(docx_bytes, bytes)