"""
Módulo para convertir HTML a DOCX.

Este módulo NO tiene dependencias de AWS ni S3.
Retorna bytes que pueden ser guardados donde sea necesario.
"""

from htmldocx import HtmlToDocx
from io import BytesIO
from docx import Document
from typing import Optional


def convert(html: str, custom_styles: Optional[dict] = None) -> bytes:
    """
    Convierte HTML a archivo DOCX (formato Word).

    Args:
        html: String con contenido HTML
        custom_styles: Diccionario con estilos personalizados (opcional)

    Returns:
        bytes: Contenido del archivo DOCX en memoria

    Raises:
        ValueError: Si el HTML está vacío
        Exception: Si hay error en la conversión

    Examples:
        >>> docx_bytes = convert("<h1>Title</h1><p>Content</p>")
        >>> type(docx_bytes)
        <class 'bytes'>
    """
    if not isinstance(html, str):
        raise TypeError("El contenido debe ser un string")
    if not html or html.strip() == "":
        raise ValueError("El contenido HTML no puede estar vacío")

    try:
        # Word vacío
        document = Document()

        parser = HtmlToDocx()

        # HTML a DOCX
        parser.add_html_to_document(html, document)

        if custom_styles:
            _apply_custom_styles(document, custom_styles)

        # Save en memoria (no en disco)
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)

        return buffer.read()

    except Exception as e:
        raise Exception(f"Error al convertir HTML a DOCX: {str(e)}")


def convert_with_template(html: str, template_path: Optional[str] = None) -> bytes:
    """
    Convierte HTML a DOCX usando una plantilla base.

    Útil para mantener formato corporativo consistente.

    Args:
        html: String con contenido HTML
        template_path: Ruta a archivo DOCX plantilla (opcional)

    Returns:
        bytes: Contenido del archivo DOCX
    """
    if template_path:
        document = Document(template_path)
    else:
        document = Document()

    parser = HtmlToDocx()
    parser.add_html_to_document(html, document)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return buffer.read()


def _apply_custom_styles(document: Document, styles: dict) -> None:
    """
    Aplica estilos personalizados al documento.

    Args:
        document: Objeto Document de python-docx
        styles: Diccionario con configuraciones de estilo
    """

    if "font_size" in styles:
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                run.font.size = styles["font_size"]